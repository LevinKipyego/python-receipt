from docx import Document
from docx.shared import Inches
import barcode
from barcode.writer import ImageWriter

def save_receipt_with_barcode():
    receipt = []
    receipt.append("sample text")
    receipt.append('_' * 40)
    receipt.append("\t\t\tBILL RECEIPT")
    receipt.append('_' * 40)
         
    doc = Document()
    # Generate the barcode
    paragraph = doc.add_paragraph('\n'.join(receipt))
    barcode_value = "1234567890"
    barcode_image = barcode.Code128(barcode_value, writer=ImageWriter()).render()

    # Save the barcode image to a file
    barcode_image_filename = "barcode.png"
    barcode_image.save(barcode_image_filename)

    # Add the barcode image to the document
    
    doc.add_picture(barcode_image_filename, width=Inches(3), height=Inches(1))

    # Save the document
    doc.save("receipt.docx")

    print("Receipt with barcode saved successfully!")

# Call the function to save the receipt with barcode
save_receipt_with_barcode()
