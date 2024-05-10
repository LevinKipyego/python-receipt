#..............imports..............#
from datetime import datetime
from prettytable import PrettyTable
import tkinter as tk
from tkinter import messagebox
from docx import Document
import pyqrcode
from tkinter import filedialog
from tkinter import font
import win32print
import win32api
from tkinter import ttk
import itemValues
from docx.shared import Inches
import barcode
from barcode.writer import ImageWriter
from docx.shared import Pt, Inches


# Define ANSI escape sequences for formatting
bold = "\033[1m"
underline = "\033[4m"
red = "\033[91m"
green = "\033[92m"
blue = "\033[94m"
end_format = "\033[0m"

#TKINTER#
window = tk.Tk()
#window.geometry("800x500")
window.title("bill receipt")

##......................FONTS.......................##

curlz = font.Font(family="Curlz Mt", size=18, weight="bold")
CurNew = font.Font(family="Courier New", size=10, weight="bold")
FootLight= font.Font(family="Footlight MT", size=10, weight="bold")
gabriola = font.Font(family="Gabriola", size=10, weight="bold")
gigi = font.Font(family="gigi", size=10, weight="bold", slant="italic")
gaudy = font.Font(family="Goudy Old Style", size=10, weight="bold")

frame = tk.Frame(window)
frame.pack()

frame_lable = tk.LabelFrame(frame,text="elgon link tech",font=curlz)
frame_lable.grid(row=0,column=0,padx=25)


def save_receipt():
    items = []
    

    def add_item():
        try:
            item = item_entry.get()
            qty = float(qty_entry.get())
            each = float(each_entry.get())
            
            total = qty * each
            if item=="":
                messagebox.showinfo("ITEM","YOU CANNOT ADD AN EMPTY ITEM!!")
            else:
                messagebox.showinfo("SUCCESS","ITEM HAS BEEN SUCCESSFULY ADDED")

        except:
            
            messagebox.showerror("","you cannot add an empty quantity or price")

        
        itm = {'Item': item, 'Quantity': qty, 'Each': each, 'Total': total}
        items.append(itm)
        item_entry.delete(0, tk.END)
        qty_entry.delete(0, tk.END)
        each_entry.delete(0, tk.END)

    def create_receipt():
        #.....LISTS....#
        
        receipt = []

        #.............initialize Document..............#
        doc = Document()

        receipt.append('_' * 40)
        receipt.append("\t\tBILL RECEIPT")
        receipt.append('_' * 40)
        receipt.append("\t\tELGON LINK TECH")
        receipt.append("\tPO BOX 1078, KITALE, KENYA")
        receipt.append("\tMobile: +254713685943,")
        receipt.append("\tEmail:cyberpunk123@gmail.com")
        receipt.append("\tDate:"+str(datetime.now())+"\n")
        receipt.append("\t\tSALES && SERVICES\n")
        receipt.append('*/*' * 14)
        
        receipt.append("\t\tSERVER:CYBERPUNK")
        receipt.append('*/*' * 14)
        receipt.append('=' * 40)
        receipt.append("\t\tPARTICULARS:")
        receipt.append('=' * 40)

        
    
        '''
        with open('receipt.doc', 'w') as file:
            file.write('\n'.join(receipt))
        '''

       #..............ADDING PARAGRAPH.....................#

        paragraph = doc.add_paragraph()
        run = paragraph.add_run('\n'.join(receipt))

        

        '''paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Inches(0)
        '''
        
        #..............ADDING STYLES.....................#

        run.font.name = 'Candara Light'
        run.font.size = Pt(10.5)

       # run.bold = True
        #run.font.italic=True
        #run.font.underline
        #run.font.underline = True
        #run.font.underline = WD_UNDERLINE.DOT_DASH

        def pretty_table():
            pretty = []

            table = PrettyTable()
            table.field_names = ["Item", "Quantity", "@", "Total"]

            for item in items:
                table.add_row([item["Item"], item["Quantity"], item["Each"], item["Total"]])
            #table.add_row(['', '', 'Person Served:', name_entry.get()])
    
            
            pretty.append(str(table))

            total_amount = sum(item['Total'] for item in items)
            pretty.append("TOTAL:"+"."*25+str(total_amount))
            name = str(name_entry.get())
            pretty.append(name)

            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Inches(0.5)

            run = paragraph.add_run('\n'.join(pretty))
            run.font.name = 'Candara Light'
            run.font.size = Pt(10.5)

        pretty_table()    

        #.........CREATE A QR CODE...................#
        def generate_qrcode(data, filename):
            qr = pyqrcode.create(data)
            qr.png(filename, scale=2)

            # Usage example
        generate_qrcode("THANKYOU FOR TRUSTING US!", "qrc.png")
        
        def save_qr_to_docx(filename):
            foter = []
            # Create a Word document
            #doc = Document()
           
            # Add barcode image to the document
            doc.add_picture(filename)
            
            foter.append('_' * 40)
            foter.append("\t\tBILL RECEIPT")
            foter.append('_' * 40)

            paragraph = doc.add_paragraph()
            
            run = paragraph.add_run("\n".join(foter))
            run.font.name = 'Candara Light'
            run.font.size = Pt(10.5)
            # Save the document
            doc.save('receipt.doc')
        save_qr_to_docx("qrc.png")

       

        '''

    #messagebox.showinfo("Receipt Saved", "Receipt saved successfully!")
    '''



    # Create the item entry and label
    item_label = tk.Label(frame_lable, text="Item:",font=CurNew,fg='blue')
    item_label.grid(row=0,column=0,sticky='w')
    item_entry = ttk.Combobox(frame_lable,width=50,values=itemValues.list)
    item_entry.grid(padx=5,pady=10,row=0,column=1)

    # Create the quantity entry and label
    qty_label = tk.Label(frame_lable, text="Quantity:",font=CurNew,fg='blue')
    qty_label.grid(row=1,column=0,sticky='w')
    qty_entry = tk.Spinbox(frame_lable,width=51,from_=0.5 ,to=1000,
                           borderwidth=0,highlightthickness=1,
                           highlightbackground="gray",
                           highlightcolor="gray")
    qty_entry.grid(padx=5,pady=10,row=1,column=1)

    # Create the each entry and label
    each_label = tk.Label(frame_lable, text="Price:",font=CurNew,fg='orange')
    each_label.grid(row=2, column=0,sticky='w')
    each_entry = ttk.Combobox(frame_lable,width=50,values=[0.5,5,10,30,50,100,150,200,250])
    each_entry.grid(padx=5,pady=10,row=2,column=1)

    # Create the customer entry and label
    name_label = tk.Label(frame_lable, text="Customer:",font=CurNew,fg='orange')
    name_label.grid(row=3, column=0,sticky='w')
    name_entry = tk.Entry(frame_lable,width=53,borderwidth=0,highlightthickness=1,highlightbackground="gray",highlightcolor="gray")
    name_entry.grid(padx=5,pady=10,row=3,column=1)

    #............FRAME 2.................#
    
    frame_lable2 = tk.LabelFrame(frame,text='frame 2')
    frame_lable2.grid(row=1, column=0,padx=25,pady=10,sticky='news')

    payment_lable = tk.Label(frame_lable2,text="PAYMENT: ")
    payment_lable.grid(row=0,column=0,padx=0,pady=10)

    cash_lable = tk.Label(frame_lable2,text="CASH/MOBILE")
    cash_lable.grid(row=0,column=1,padx=0,pady=10)

    paid_lable = tk.Label(frame_lable2,text="PAID: ")
    paid_lable.grid(row=0,column=2,padx=5,pady=10)

    Mpaid_lable = tk.Label(frame_lable2,text="100")
    Mpaid_lable.grid(row=0,column=3,padx=0,pady=10)

    total_lable = tk.Label(frame_lable2,text="TOTAL: ")
    total_lable.grid(row=0,column=4,padx=0,pady=10)

    tota_lable = tk.Label(frame_lable2,text="500")
    tota_lable.grid(row=0,column=5,padx=5,pady=10)

    change_lable = tk.Label(frame_lable2,text="CHANGE")
    change_lable.grid(row=0,column=6,padx=5,pady=10)

    change = tk.Label(frame_lable2,text="100")
    change.grid(row=0,column=7,padx=5,pady=10)

   # cash_lable = tk.Label(frame_lable2,text="CASH")
    #cash_lable.grid(row=0,column=0,padx=5,pady=10)

    # Create the add item button
    add_button = tk.Button(frame_lable, text="Add Item", command=add_item,
                            
                            bg='lightgreen',
                            height=1,width=11,
                            borderwidth=0,highlightthickness=1,
                            highlightbackground="gray",highlightcolor="gray")
    add_button.grid(padx=5,pady=10,row=0,column=2,sticky='w')

    # Create the create receipt button
    create_button = tk.Button(frame_lable, text="Create Receipt", 
                                
                                command=create_receipt,bg='lightgreen',
                                height=1,width=11,
                                borderwidth=0,highlightthickness=1,
                                highlightbackground="gray",highlightcolor="gray")
    create_button.grid(padx=5,pady=10,row=1,column=2,sticky='w')
    
    
save_receipt()

def open_file_dialog():
    filepath = filedialog.askopenfilename()
    if filepath:
        print_file(filepath)

def print_file(filepath):
    printer_name = win32print.GetDefaultPrinter()
    if not printer_name:
        print("No default printer found.")
        return 

    try:
        win32api.ShellExecute(0, "print", filepath, '/d:"%s"' % printer_name, ".", 0)
        print("File sent to printer successfully.")
    except Exception as e:
        print("Failed to print file:", e)



# Create a button to open the file dialog
button = tk.Button(frame_lable, text="Open File",bg='lightgreen',
                   command=open_file_dialog,height=1,width=11,
                   borderwidth=0,highlightthickness=1
                   ,highlightbackground="gray",highlightcolor="gray")
button.grid(padx=5,pady=10,row=2,column=2,sticky='w')

#paddings
for widget in frame_lable.winfo_children():
    widget.grid_configure(padx=5,pady=10)
    
# Run the main window event loop
window.mainloop()
